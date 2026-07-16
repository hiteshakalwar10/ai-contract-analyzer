import os
import json
import fitz
import streamlit as st
from datetime import datetime
from langchain_groq import ChatGroq
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_core.documents import Document
from docx import Document as DocxDocument
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from pydantic import BaseModel, Field
from typing import List

os.makedirs("uploads", exist_ok=True)
os.makedirs("outputs", exist_ok=True)
os.makedirs("vectorstore", exist_ok=True)

GROQ_API_KEY = os.environ.get("GROQ_API_KEY", "")

st.set_page_config(page_title="AI Contract Analyzer", layout="wide")

@st.cache_resource
def get_llm():
    return ChatGroq(groq_api_key=GROQ_API_KEY, model_name="llama-3.3-70b-versatile", temperature=0)

@st.cache_resource
def get_embeddings():
    return HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")

def extract_pdf_text(pdf_path):
    text = ""
    doc = fitz.open(pdf_path)
    for page in doc:
        text += page.get_text()
    doc.close()
    return text

def chunk_text(text):
    splitter = RecursiveCharacterTextSplitter(chunk_size=1200, chunk_overlap=200)
    chunks = splitter.split_text(text)
    return [Document(page_content=c) for c in chunks]

def build_vectorstore(docs):
    embeddings = get_embeddings()
    vectorstore = FAISS.from_documents(docs, embeddings)
    vectorstore.save_local("vectorstore/contract_index")
    return vectorstore

def load_vectorstore():
    embeddings = get_embeddings()
    return FAISS.load_local("vectorstore/contract_index", embeddings, allow_dangerous_deserialization=True)

def run_prompt(prompt):
    llm = get_llm()
    response = llm.invoke(prompt)
    return response.content

def analyze_contract(full_text):
    prompt = f"""
You are a senior legal contract analyst. Analyze the contract text below and return ONLY valid JSON with this exact schema, no markdown, no extra text:

{{
"executive_summary": "string",
"parties_involved": ["string"],
"important_dates": ["string"],
"payment_terms": "string",
"key_clauses": ["string"],
"obligations": ["string"],
"risk_analysis": ["string"]
}}

Contract text:
{full_text[:12000]}
"""
    raw = run_prompt(prompt)
    raw = raw.strip()
    if raw.startswith("```"):
        raw = raw.split("```")[1]
        if raw.startswith("json"):
            raw = raw[4:]
    try:
        data = json.loads(raw)
    except Exception:
        data = {
            "executive_summary": raw,
            "parties_involved": [],
            "important_dates": [],
            "payment_terms": "",
            "key_clauses": [],
            "obligations": [],
            "risk_analysis": []
        }
    return data

def answer_question(vectorstore, question):
    retriever = vectorstore.as_retriever(search_kwargs={"k": 4})
    relevant_docs = retriever.invoke(question)
    context = "\n\n".join([d.page_content for d in relevant_docs])
    prompt = f"""
Answer the question using ONLY the context below from the contract. If the answer is not present, say "This information is not present in the contract."

Context:
{context}

Question: {question}
"""
    return run_prompt(prompt)

def generate_pdf_report(data, output_path):
    doc = SimpleDocTemplate(output_path, pagesize=A4)
    styles = getSampleStyleSheet()
    story = []
    story.append(Paragraph("Contract Analysis Report", styles["Title"]))
    story.append(Spacer(1, 12))

    sections = [
        ("Executive Summary", data.get("executive_summary", "")),
        ("Parties Involved", ", ".join(data.get("parties_involved", []))),
        ("Important Dates", ", ".join(data.get("important_dates", []))),
        ("Payment Terms", data.get("payment_terms", "")),
        ("Key Clauses", ", ".join(data.get("key_clauses", []))),
        ("Obligations", ", ".join(data.get("obligations", []))),
        ("Risk Analysis", ", ".join(data.get("risk_analysis", [])))
    ]

    for title, content in sections:
        story.append(Paragraph(title, styles["Heading2"]))
        story.append(Paragraph(content if content else "N/A", styles["Normal"]))
        story.append(Spacer(1, 10))

    doc.build(story)
    return output_path

def generate_docx_report(data, output_path):
    doc = DocxDocument()
    doc.add_heading("Contract Analysis Report", 0)

    sections = [
        ("Executive Summary", data.get("executive_summary", "")),
        ("Parties Involved", ", ".join(data.get("parties_involved", []))),
        ("Important Dates", ", ".join(data.get("important_dates", []))),
        ("Payment Terms", data.get("payment_terms", "")),
        ("Key Clauses", ", ".join(data.get("key_clauses", []))),
        ("Obligations", ", ".join(data.get("obligations", []))),
        ("Risk Analysis", ", ".join(data.get("risk_analysis", [])))
    ]

    for title, content in sections:
        doc.add_heading(title, level=1)
        doc.add_paragraph(content if content else "N/A")

    doc.save(output_path)
    return output_path

if "analysis_data" not in st.session_state:
    st.session_state.analysis_data = None
if "vectorstore" not in st.session_state:
    st.session_state.vectorstore = None
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []
if "pdf_path" not in st.session_state:
    st.session_state.pdf_path = None

with st.sidebar:
    st.title("📑 AI Contract Analyzer")
    uploaded_file = st.file_uploader("Upload Contract PDF", type=["pdf"])

    if uploaded_file is not None:
        pdf_path = os.path.join("uploads", uploaded_file.name)
        with open(pdf_path, "wb") as f:
            f.write(uploaded_file.getbuffer())
        st.session_state.pdf_path = pdf_path
        st.success("File uploaded successfully")

    analyze_btn = st.button("🔍 Analyze Contract", use_container_width=True)

st.markdown("## AI Contract Analyzer")
st.caption("Upload a contract, get structured analysis, and chat with it using RAG.")

if analyze_btn:
    if st.session_state.pdf_path is None:
        st.warning("Please upload a PDF first.")
    else:
        with st.spinner("Extracting text..."):
            full_text = extract_pdf_text(st.session_state.pdf_path)

        with st.spinner("Chunking and building vector store..."):
            docs = chunk_text(full_text)
            vectorstore = build_vectorstore(docs)
            st.session_state.vectorstore = vectorstore

        with st.spinner("Running structured analysis..."):
            data = analyze_contract(full_text)
            st.session_state.analysis_data = data

        st.success("Analysis complete")

tab1, tab2, tab3 = st.tabs(["📊 Analysis", "💬 Chat with Contract", "⬇️ Download Reports"])

with tab1:
    data = st.session_state.analysis_data
    if data is None:
        st.info("Upload and analyze a contract to see results.")
    else:
        st.subheader("Executive Summary")
        st.write(data.get("executive_summary", ""))

        col1, col2 = st.columns(2)
        with col1:
            st.subheader("Parties Involved")
            for p in data.get("parties_involved", []):
                st.write(f"- {p}")

            st.subheader("Important Dates")
            for d in data.get("important_dates", []):
                st.write(f"- {d}")

            st.subheader("Payment Terms")
            st.write(data.get("payment_terms", ""))

        with col2:
            st.subheader("Key Clauses")
            for c in data.get("key_clauses", []):
                st.write(f"- {c}")

            st.subheader("Obligations")
            for o in data.get("obligations", []):
                st.write(f"- {o}")

            st.subheader("Risk Analysis")
            for r in data.get("risk_analysis", []):
                st.write(f"- ⚠️ {r}")

with tab2:
    if st.session_state.vectorstore is None:
        st.info("Analyze a contract first to enable chat.")
    else:
        for role, msg in st.session_state.chat_history:
            with st.chat_message(role):
                st.write(msg)

        user_question = st.chat_input("Ask something about the contract...")
        if user_question:
            st.session_state.chat_history.append(("user", user_question))
            with st.spinner("Thinking..."):
                answer = answer_question(st.session_state.vectorstore, user_question)
            st.session_state.chat_history.append(("assistant", answer))
            st.rerun()

with tab3:
    if st.session_state.analysis_data is None:
        st.info("Analyze a contract first to generate reports.")
    else:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        pdf_output = f"outputs/contract_report_{timestamp}.pdf"
        docx_output = f"outputs/contract_report_{timestamp}.docx"

        if st.button("Generate PDF Report"):
            generate_pdf_report(st.session_state.analysis_data, pdf_output)
            with open(pdf_output, "rb") as f:
                st.download_button("⬇️ Download PDF", f, file_name="contract_report.pdf")

        if st.button("Generate DOCX Report"):
            generate_docx_report(st.session_state.analysis_data, docx_output)
            with open(docx_output, "rb") as f:
                st.download_button("⬇️ Download DOCX", f, file_name="contract_report.docx")
